"""Generate OMLT Master Presentation DOCX — bilingual, brand-aligned."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

# Palette
TERRA = RGBColor(0xC8, 0x44, 0x2C)
CREAM = RGBColor(0xF6, 0xEC, 0xD6)
NAVY  = RGBColor(0x1B, 0x22, 0x35)
YOLK  = RGBColor(0xE6, 0xA5, 0x32)
GRAY  = RGBColor(0x6B, 0x77, 0x75)

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width  = Cm(29.7)   # A4 landscape
section.page_height = Cm(21.0)
section.left_margin = Cm(2.4)
section.right_margin = Cm(2.4)
section.top_margin = Cm(2.2)
section.bottom_margin = Cm(2.2)

# Default style
style = doc.styles["Normal"]
style.font.name = "Tajawal"
style.font.size = Pt(11)

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)

def shade_cell(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)

def add_para(text, bold=False, italic=False, size=11, color=None, font="Tajawal",
             align=WD_ALIGN_PARAGRAPH.LEFT, rtl=True, space_after=4, space_before=0):
    p = doc.add_paragraph()
    p.alignment = align
    if rtl: set_rtl(p)
    pf = p.paragraph_format
    pf.space_after = Pt(space_after)
    pf.space_before = Pt(space_before)
    pf.line_spacing = 1.4
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    return p

def add_kicker(text):
    add_para(text.upper(), bold=True, size=9, color=TERRA, font="Space Grotesk",
             align=WD_ALIGN_PARAGRAPH.LEFT, rtl=False, space_after=4)

def add_h1(text, ar_text=None, color=NAVY):
    add_para(text, bold=True, size=32, color=color, font="Frank Ruhl Libre",
             align=WD_ALIGN_PARAGRAPH.LEFT, rtl=False, space_after=2, space_before=2)
    if ar_text:
        add_para(ar_text, bold=True, size=22, color=color, font="Tajawal",
                 align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=10)

def add_h2(text, color=NAVY):
    add_para(text, bold=True, size=20, color=color, font="Frank Ruhl Libre",
             align=WD_ALIGN_PARAGRAPH.LEFT, rtl=False, space_after=8, space_before=6)

def add_h3(text, color=NAVY):
    add_para(text, bold=True, size=14, color=color, font="Frank Ruhl Libre",
             align=WD_ALIGN_PARAGRAPH.LEFT, rtl=False, space_after=4, space_before=4)

def add_body_ar(text, size=11, color=None):
    add_para(text, size=size, color=color or NAVY, font="Tajawal",
             align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)

def add_body_en(text, size=11, italic=False, color=None):
    add_para(text, size=size, italic=italic, color=color or NAVY, font="Frank Ruhl Libre",
             align=WD_ALIGN_PARAGRAPH.LEFT, rtl=False, space_after=4)

def add_bullet_ar(text):
    p = doc.add_paragraph(style="List Bullet")
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = "Tajawal"
    run.font.size = Pt(11)
    run.font.color.rgb = NAVY

def add_divider():
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "C8442C")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)

def page_break():
    doc.add_page_break()

def page_label(num, total, name_en, name_ar):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_after = Pt(2)
    run = p.add_run(f"{num:03d} / {total:03d}   ·   ")
    run.font.name = "Space Grotesk"; run.font.size = Pt(8); run.font.color.rgb = TERRA; run.bold = True
    run = p.add_run(f"{name_en}   —   ")
    run.font.name = "Space Grotesk"; run.font.size = Pt(8); run.font.color.rgb = GRAY
    run = p.add_run(name_ar)
    run.font.name = "Tajawal"; run.font.size = Pt(9); run.font.color.rgb = GRAY
    add_divider()

# ─────────────────────────────────────────────────────────────
# COVER (Page 1)
# ─────────────────────────────────────────────────────────────
# Title block
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("OMLT · MASTER PRESENTATION · V 1.0")
run.font.name = "Space Grotesk"; run.font.size = Pt(10); run.font.color.rgb = TERRA; run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(20)
run = p.add_run("ABHA · KSA")
run.font.name = "Space Grotesk"; run.font.size = Pt(9); run.font.color.rgb = GRAY

# Editorial cover — type-led, no logo. Logo gets its dedicated hero on page 15.
# Add vertical breathing room before the title
for _ in range(3):
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)

# Kicker
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A  BREAKFAST  KITCHEN")
run.font.name = "Space Grotesk"; run.font.size = Pt(11); run.bold = True; run.font.color.rgb = YOLK
p.paragraph_format.space_after = Pt(12)

# Title
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A clear, simple morning.")
run.font.name = "Frank Ruhl Libre"; run.font.size = Pt(56); run.bold = True; run.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
run = p.add_run("بداية صباح واضحة وبسيطة.")
run.font.name = "Tajawal"; run.font.size = Pt(20); run.bold = True; run.font.color.rgb = TERRA
p.paragraph_format.space_after = Pt(18)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_rtl(p)
run = p.add_run("عرض مرجعي معتمد لعلامة OMLT — تعريف المشروع، الرؤية، التكاليف، والهوية. وثيقة كاملة جاهزة للعرض على المستثمر والشركاء التنفيذيين.")
run.font.name = "Tajawal"; run.font.size = Pt(11); run.font.color.rgb = GRAY

page_break()

# ─────────────────────────────────────────────────────────────
# CONTENTS (Page 2)
# ─────────────────────────────────────────────────────────────
page_label(2, 26, "Contents", "الفهرس")
add_kicker("The map")
add_h1("Twenty-six pages, one egg.", "ست وعشرون صفحة، بيضة واحدة.")

toc = [
    ("01", "OMLT Concept", "مفهوم المشروع", "003"),
    ("·",  "Location & Space", "الموقع والفراغ", "007"),
    ("02", "Vision & Purpose", "الرؤية والهدف", "008"),
    ("03", "Business Setup & Costs", "التكاليف", "011"),
    ("04", "Approved Logo", "الشعار المعتمد", "016"),
    ("05", "Brand Colors", "الألوان", "019"),
    ("06", "Typography", "الخطوط", "021"),
    ("07", "Visual Identity", "الهوية البصرية", "023"),
    ("08", "Sensory Identity", "الهوية الحسّية", "024"),
    ("09", "Audio Identity", "الهوية الصوتية", "025"),
]
tbl = doc.add_table(rows=len(toc), cols=4)
tbl.autofit = False
widths = [Cm(1.2), Cm(9), Cm(7), Cm(2)]
for row in tbl.rows:
    for i, w in enumerate(widths):
        row.cells[i].width = w
for i, (n, t_en, t_ar, p_) in enumerate(toc):
    row = tbl.rows[i]
    # number
    c = row.cells[0].paragraphs[0]; c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = c.add_run(n); r.font.name = "Space Grotesk"; r.font.size = Pt(10); r.bold = True; r.font.color.rgb = TERRA
    # english
    c = row.cells[1].paragraphs[0]; c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = c.add_run(t_en); r.font.name = "Frank Ruhl Libre"; r.font.size = Pt(13); r.font.color.rgb = NAVY
    # arabic
    c = row.cells[2].paragraphs[0]; c.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_rtl(c)
    r = c.add_run(t_ar); r.font.name = "Tajawal"; r.font.size = Pt(11); r.font.color.rgb = GRAY
    # page
    c = row.cells[3].paragraphs[0]; c.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = c.add_run(p_); r.font.name = "Space Grotesk"; r.font.size = Pt(9); r.font.color.rgb = GRAY

page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 3 · CONCEPT — DEFINITION
# ─────────────────────────────────────────────────────────────
page_label(3, 26, "Concept · Definition", "تعريف المشروع")
add_kicker("The concept")
add_h1("A specialty breakfast kitchen.", "مطبخ فطور متخصّص.")
add_body_ar("OMLT مطبخ فطور متخصّص يعتمد على الأومليت كمنتج رئيسي ضمن قائمة قصيرة ومحدّدة، يُقدّم تجربة صباحية واضحة وسريعة داخل مساحة صغيرة ذات مطبخ مفتوح.")
add_body_en("OMLT is a specialty breakfast kitchen with the omelette as the anchor dish, served from a short and focused menu in a small space with an open kitchen.", italic=True, color=TERRA)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 4 · CONCEPT — OPERATIONAL CHARACTERISTICS
# ─────────────────────────────────────────────────────────────
page_label(4, 26, "Concept · Operational characteristics", "الخصائص التشغيلية")
add_kicker("How it operates")
add_h2("قائمة قصيرة · إيقاع صباحي ثابت · مطبخ مفتوح.")

bullets = [
    "قائمة مختصرة من 5–7 أصناف أساسية",
    "تشغيل صباحي محدّد — يفتح 7:00 ص ويُغلق 3:00 م",
    "إيقاع تشغيل ثابت وواضح طوال أيام العمل",
    "تجربة داخلية مركّزة على الفطور فقط",
    "مطبخ مفتوح ظاهر بالكامل أمام العميل",
    "جلسات محدودة — 8 إلى 10 طاولات",
    "سرعة تقديم ودوران عملاء مناسب لفترة الصباح",
    "بناء قاعدة عملاء متكرّرين ضمن الروتين الصباحي",
]
for b in bullets: add_bullet_ar(b)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 5 · CONCEPT — WHAT IT IS NOT
# ─────────────────────────────────────────────────────────────
page_label(5, 26, "Concept · What it is not", "ما لا يكون")
add_kicker("Discipline by exclusion")
add_h1("ما لا يكون OMLT.", "What OMLT is not.")
add_body_ar("قوة المشروع تأتي من حدوده الواضحة. هذه القائمة تحمي الفكرة من التمدّد وتُبقي التشغيل بسيطًا.")

nots = [
    ("ليس مطعم برانش", "Not brunch"),
    ("ليس مقهى تقليديًّا", "Not a traditional café"),
    ("ليس مطعمًا طوال اليوم", "Not all-day dining"),
    ("ليست تجربة فاخرة مبالغًا بها", "Not over-luxury"),
    ("ليست قائمة كبيرة متعدّدة", "Not a long menu"),
]
tbl = doc.add_table(rows=len(nots), cols=2)
tbl.autofit = False
for row in tbl.rows:
    row.cells[0].width = Cm(11); row.cells[1].width = Cm(7)
for i, (ar, en) in enumerate(nots):
    row = tbl.rows[i]
    c = row.cells[0].paragraphs[0]; c.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(c)
    r = c.add_run("✕  " + ar); r.font.name = "Tajawal"; r.font.size = Pt(13); r.bold = True; r.font.color.rgb = NAVY
    c = row.cells[1].paragraphs[0]; c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = c.add_run(en); r.font.name = "Space Grotesk"; r.font.size = Pt(10); r.font.color.rgb = GRAY
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 6 · CONCEPT — OPEN KITCHEN
# ─────────────────────────────────────────────────────────────
page_label(6, 26, "Concept · Open kitchen", "المطبخ المفتوح")
add_kicker("Heart of the experience")
add_h1("The open kitchen is the show.", "المطبخ المفتوح قلب التجربة.")
add_body_ar("المطبخ المفتوح ليس عنصر تصميم — بل مركز التجربة اليومية. يظهر مباشرة عند الدخول، فيُشاهد العميل التحضير والطهي والتشغيل، وتصبح حركة العمل جزءًا من الهوية.")
for b in ["صوت البيض وحركة المقلاة", "تحضير الأومليت أمام العميل", "رائحة الخبز الطازج والقهوة", "التشغيل الصباحي ظاهر للعميل"]:
    add_bullet_ar(b)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 7 · CONCEPT — LOCATION & SPACE  (new)
# ─────────────────────────────────────────────────────────────
page_label(7, 26, "Concept · Location & Space", "الموقع والفراغ")
add_kicker("Where the morning lives")
add_h1("A street, a sunrise, a room that breathes.", "موقع، ضوء، وفراغ يتنفّس.")
add_body_ar("الموقع ليس مجرّد عنوان — هو جزء من تجربة OMLT. الاتجاه، الحيّ، وحجم الفراغ كلها قرارات تخدم الإيقاع الصباحي وشعور المكان.")

loc_items = [
    ("01 · Orientation",  "اتجاه شرقي",
     "تفضيل الموقع المُتجه شرقًا قدر الإمكان — ضوء الصباح الطبيعي جزء من تجربة OMLT، يدعم هدوء الجو الصباحي."),
    ("02 · Neighborhood", "حيّ نابض بالحياة",
     "شارع حيوي ومأهول في أبها — لا منعزل، لا مخفي. ظاهر للناس وضمن إيقاع يومهم الطبيعي."),
    ("03 · Footprint",    "٨ – ١٠ طاولات",
     "فراغ يستوعب ٨–١٠ طاولات بمسافات مريحة — صغير ليبقى شخصيًّا، كبير ليبقى حيًّا."),
]
for label, h, body in loc_items:
    add_para(label, bold=True, size=10, color=TERRA, font="Space Grotesk", rtl=False, space_after=2, space_before=8)
    add_para(h, bold=True, size=14, color=NAVY, font="Tajawal", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_body_ar(body, size=11)

add_divider()
add_para("THE FEELING WE TEST FOR", bold=True, size=10, color=TERRA, font="Space Grotesk", rtl=False, space_before=6, space_after=2)
add_para("Small enough to feel personal, large enough to feel alive.",
         bold=True, italic=True, size=14, color=NAVY, font="Frank Ruhl Libre",
         align=WD_ALIGN_PARAGRAPH.LEFT, rtl=False, space_after=2)
add_body_ar("صغير ليُحسّ، كبير ليحيا.", size=12)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 7 · VISION — MORNING ROUTINE
# ─────────────────────────────────────────────────────────────
page_label(8, 26, "Vision · Become part of the morning", "جزء من الصباح")
add_kicker("Vision")
add_h1("Become part of their morning.", "أن نكون جزءًا من صباحهم.")
add_body_ar("نهدف إلى صناعة مكان يعود إليه العميل ضمن روتينه الصباحي — لا يأتي مرّة واحدة، بل يجعله جزءًا من أيّامه. النجاح يُقاس بالعودة المتكرّرة، لا بعدد الزيارات الأولى.")
add_body_en("We aim to build a place the customer returns to as part of their morning routine.", italic=True, color=TERRA)

add_divider()
add_h3("الهدف الأساسي · Core purpose")
add_body_ar("OMLT يُقدّم وجبة فطور بسيطة ومُتقَنة، في مكان واضح ومألوف، بإيقاع صباحي ثابت. الهدف ليس الإثارة أو التميّز اللحظي، بل بناء عادة يومية يمكن الاعتماد عليها.")
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 8 · VISION — PRINCIPLES
# ─────────────────────────────────────────────────────────────
page_label(9, 26, "Vision · Operational principles", "المبادئ التشغيلية")
add_kicker("Four principles, daily")
add_h2("أربعة مبادئ تحكم كل قرار في OMLT.")

principles = [
    ("01 · Plain done well", "البساطة في التنفيذ", "نُتقن قائمة محدودة بدلًا من توسيع المنيو."),
    ("02 · Clear by design", "الوضوح في التجربة", "كل قرار يخدم سهولة الزيارة من الدخول للمغادرة."),
    ("03 · Steady quality", "ثبات الجودة", "نفس الإيقاع والجودة كل صباح، طوال أيام العمل."),
    ("04 · Build the habit", "بناء العادة", "كل قرار يُختبَر: هل يدعم عودة العميل في الزيارة التالية؟"),
]
for label, h, body in principles:
    add_para(label, bold=True, size=10, color=TERRA, font="Space Grotesk", rtl=False, space_after=2, space_before=8)
    add_para(h, bold=True, size=14, color=NAVY, font="Tajawal", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_body_ar(body, size=11)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 9 · VISION — EXPANSION ROADMAP
# ─────────────────────────────────────────────────────────────
page_label(10, 26, "Vision · Realistic expansion", "خطة التوسّع")
add_kicker("A three-phase path")
add_h2("نموذج واحد يُثبت أولًا، ثم يُدرَس التوسّع.")

phases = [
    ("Phase 01 — Prove the idea", "المرحلة الأولى — إثبات الفكرة",
     ["افتتاح الفرع الأول", "بناء قاعدة عملاء أساسية", "تحسين التشغيل اليومي", "إثبات نجاح الفكرة عمليًّا"]),
    ("Phase 02 — Document & develop", "المرحلة الثانية — توثيق وتطوير",
     ["توثيق الوصفات بشكل قياسي", "تطوير إجراءات تشغيل موحّدة", "تحسين التجربة بناءً على الواقع", "الاستقرار المالي للفرع الأول"]),
    ("Phase 03 — Study expansion", "المرحلة الثالثة — دراسة التوسّع",
     ["دراسة فرع ثانٍ عند نجاح النموذج", "تقييم الموقع بناءً على البيانات", "الحفاظ على نفس النموذج والهوية"]),
]
for en, ar, items in phases:
    add_para(en, bold=True, size=10, color=TERRA, font="Space Grotesk", rtl=False, space_after=2, space_before=8)
    add_para(ar, bold=True, size=13, color=NAVY, font="Tajawal", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=4)
    for b in items: add_bullet_ar(b)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 10 · COSTS — STARTUP COSTS
# ─────────────────────────────────────────────────────────────
page_label(11, 26, "Business setup · Startup costs", "التكاليف الأولية")
add_kicker("Capital required")
add_h1("Setup budget, itemised.", "ميزانية التأسيس.")
add_body_ar("الأرقام تقديرية لنموذج تشغيلي صغير — مساحة محدودة، 8–10 طاولات، فريق مختصر.")

rows_data = [
    ("البند · ITEM", "التفاصيل · DETAIL", "SAR"),
    ("تهيئة وتجهيز · Interior", "Build-out · joinery · lighting", "120,000 – 180,000"),
    ("معدّات المطبخ · Equipment", "Induction · pans · refrigeration · ventilation · POS", "70,000 – 100,000"),
    ("أثاث وطاولات · Furniture", "Tables · chairs · ceramics", "30,000 – 50,000"),
    ("هوية وطباعة · Branding", "Print · signage · kits", "20,000 – 30,000"),
    ("تراخيص وتأسيس · Licensing", "CR · municipality · health · setup", "20,000 – 35,000"),
    ("سيولة تشغيل · Working capital", "First months runway", "70,000 – 100,000"),
    ("الإجمالي المتوقّع · Total", "Estimated launch range", "330,000 – 495,000"),
]
tbl = doc.add_table(rows=len(rows_data), cols=3)
tbl.style = "Light Grid Accent 1"
tbl.autofit = False
for r in tbl.rows:
    r.cells[0].width = Cm(7); r.cells[1].width = Cm(9); r.cells[2].width = Cm(5)

for i, (a, b, c) in enumerate(rows_data):
    row = tbl.rows[i]
    cells = [row.cells[0], row.cells[1], row.cells[2]]
    for cell, txt, align in zip(cells, (a, b, c), (WD_ALIGN_PARAGRAPH.RIGHT, WD_ALIGN_PARAGRAPH.LEFT, WD_ALIGN_PARAGRAPH.RIGHT)):
        p = cell.paragraphs[0]; p.alignment = align
        if align == WD_ALIGN_PARAGRAPH.RIGHT and "·" in txt and "SAR" not in txt and any(ord(ch) > 127 for ch in txt):
            set_rtl(p)
        r = p.add_run(txt)
        if i == 0:
            r.bold = True; r.font.color.rgb = TERRA; r.font.name = "Space Grotesk"; r.font.size = Pt(10)
            shade_cell(cell, "FBF6EA")
        elif i == len(rows_data)-1:
            r.bold = True; r.font.color.rgb = TERRA; r.font.size = Pt(11)
            r.font.name = "Tajawal" if any(ord(ch) > 127 for ch in txt) else "Space Grotesk"
            shade_cell(cell, "F6ECD6")
        else:
            r.font.color.rgb = NAVY; r.font.size = Pt(10.5)
            r.font.name = "Tajawal" if any(ord(ch) > 127 for ch in txt) else "Space Grotesk"
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 11 · COSTS — LAUNCH BUDGET SUMMARY
# ─────────────────────────────────────────────────────────────
page_label(12, 26, "Business setup · Launch budget", "ميزانية الإطلاق")
add_kicker("Expected launch budget")
add_h1("330K — 495K SAR.", "بين 330 و 495 ألف ريال.")

summary = [
    ("Interior · تهيئة", "36%", "تهيئة، تشطيب، إضاءة، تنفيذ المطبخ المفتوح."),
    ("Equipment · معدّات", "30%", "معدّات الطبخ والتبريد والقهوة وأدوات التشغيل."),
    ("Branding · هوية", "6%", "الهوية، الطباعة، اللافتات، التغليف الأساسي."),
    ("Reserves · سيولة", "14%", "سيولة التشغيل لأول أشهر — احتياط أمان."),
]
tbl = doc.add_table(rows=len(summary), cols=3)
tbl.autofit = False
for r in tbl.rows:
    r.cells[0].width = Cm(6); r.cells[1].width = Cm(3); r.cells[2].width = Cm(12)
for i, (name, pct, body) in enumerate(summary):
    row = tbl.rows[i]
    p = row.cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(name); r.font.name = "Space Grotesk"; r.bold = True; r.font.size = Pt(11); r.font.color.rgb = NAVY
    p = row.cells[1].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(pct); r.font.name = "Frank Ruhl Libre"; r.bold = True; r.font.size = Pt(20); r.font.color.rgb = TERRA
    p = row.cells[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p)
    r = p.add_run(body); r.font.name = "Tajawal"; r.font.size = Pt(10.5); r.font.color.rgb = NAVY

add_para("الأرقام تقديرية لمرحلة التأسيس. الأثاث والتراخيص يدخلان ضمن نسبة مشتركة قدرها 14% من إجمالي الميزانية.",
         size=10, color=GRAY, font="Tajawal", align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=12)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 12 · COSTS — STAFFING PLAN
# ─────────────────────────────────────────────────────────────
page_label(13, 26, "Business setup · Staffing plan", "فريق التشغيل")
add_kicker("A small, focused team")
add_h2("ستة أدوار، فريقٌ صغير ودقيق.")
add_body_ar("فريق مناسب لمساحة 8–10 طاولات وساعات تشغيل صباحية فقط (7:00 ص – 3:00 م).")

roles = [
    ("Head Chef", "شيف رئيسي", "مسؤول عن الجودة والإشراف والتشغيل اليومي للمطبخ."),
    ("Kitchen Assistant", "مساعد مطبخ", "التجهيز والتحضير ودعم التشغيل في وقت الذروة."),
    ("Beverage Specialist", "مسؤول مشروبات", "تحضير القهوة والشاي والمشروبات الصباحية."),
    ("Service Staff", "موظّف خدمة", "الاستقبال داخل الصالة، تقديم الطلبات، متابعة الطاولات."),
    ("Cashier", "كاشير", "نقاط البيع، التقفيل اليومي، الجرد، طلبات التوريد."),
    ("Cleaning & Support", "نظافة ومساندة", "نظافة الصالة والمطبخ ودعم تشغيلي عام."),
]
for en, ar, body in roles:
    add_para(en, bold=True, size=10, color=TERRA, font="Space Grotesk", rtl=False, space_after=0, space_before=6)
    add_para(ar, bold=True, size=12, color=NAVY, font="Tajawal", align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2)
    add_body_ar(body, size=10.5)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 13 · COSTS — ESTIMATED SALARIES
# ─────────────────────────────────────────────────────────────
page_label(14, 26, "Business setup · Estimated salaries", "الرواتب الشهرية")
add_kicker("Monthly payroll")
add_h1("A lean monthly payroll.", "رواتب شهرية مدروسة.")
add_body_ar("الأرقام شاملة لكل بنود التوظيف، مبنية على ساعات تشغيل صباحية فقط (7:00 – 3:00).")

salaries = [
    ("الوظيفة · ROLE", "ENGLISH", "SAR / month"),
    ("شيف رئيسي", "Head Chef", "5,500 – 6,500"),
    ("مساعد مطبخ", "Kitchen Assistant", "2,500 – 3,500"),
    ("مسؤول مشروبات", "Beverage Specialist", "2,500 – 3,000"),
    ("موظّف خدمة", "Service Staff", "2,500 – 3,000"),
    ("كاشير", "Cashier", "3,000 – 3,500"),
    ("نظافة ومساندة", "Cleaning & Support", "1,800 – 2,200"),
    ("إجمالي الرواتب الشهري · Total", "Total monthly payroll", "17,800 – 21,700"),
]
tbl = doc.add_table(rows=len(salaries), cols=3)
tbl.style = "Light Grid Accent 1"
tbl.autofit = False
for r in tbl.rows:
    r.cells[0].width = Cm(7); r.cells[1].width = Cm(7); r.cells[2].width = Cm(6)
for i, (a, b, c) in enumerate(salaries):
    row = tbl.rows[i]
    p = row.cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if any(ord(ch) > 127 for ch in a): set_rtl(p)
    r = p.add_run(a); r.font.name = "Tajawal"; r.font.size = Pt(10.5)
    p = row.cells[1].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(b); r.font.name = "Frank Ruhl Libre"; r.italic = True; r.font.size = Pt(10.5)
    p = row.cells[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(c); r.font.name = "Space Grotesk"; r.font.size = Pt(10.5); r.bold = True
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = NAVY
    if i == 0:
        for cell in row.cells:
            shade_cell(cell, "FBF6EA")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = TERRA; run.bold = True
    if i == len(salaries)-1:
        for cell in row.cells:
            shade_cell(cell, "F6ECD6")
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = TERRA; run.bold = True
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 14 · COSTS — REVENUE MODEL
# ─────────────────────────────────────────────────────────────
page_label(15, 26, "Business setup · Revenue model", "نموذج الإيرادات")
add_kicker("Realistic financial targets")
add_h2("إيرادات مبنية على استقرار، لا على وعود.")

rev = [
    ("Avg ticket", "35 – 55 SAR", "متوسّط الفاتورة للضيف الواحد"),
    ("Daily guests", "50 – 80", "العملاء يوميًا"),
    ("Operating days", "6 / week", "أيام العمل الأسبوعية"),
    ("Hours", "07:00 — 15:00", "ساعات التشغيل — وردية صباحية فقط"),
]
tbl = doc.add_table(rows=len(rev), cols=3)
tbl.autofit = False
for r in tbl.rows:
    r.cells[0].width = Cm(6); r.cells[1].width = Cm(6); r.cells[2].width = Cm(9)
for label, val, ar in rev:
    row = tbl.add_row() if False else tbl.rows[rev.index((label, val, ar))]
    p = row.cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(label); r.font.name = "Space Grotesk"; r.bold = True; r.font.color.rgb = TERRA; r.font.size = Pt(10)
    p = row.cells[1].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(val); r.font.name = "Frank Ruhl Libre"; r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY
    p = row.cells[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p)
    r = p.add_run(ar); r.font.name = "Tajawal"; r.font.size = Pt(11); r.font.color.rgb = NAVY

add_divider()
add_h3("مبادئ النموذج المالي · Financial principles")
for b in [
    "الهدف الأساسي بناء قاعدة عملاء متكرّرين — لا تحقيق إيرادات مرتفعة سريعًا",
    "الاستقرار التشغيلي في الفرع الأول قبل التفكير في أي توسّع",
    "التحكّم في تكاليف المواد والإنتاج عبر القائمة المختصرة",
    "الفريق المختصر يُقلّل من الرواتب الشهرية ويبسّط التشغيل",
]:
    add_bullet_ar(b)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 15 · LOGO HERO
# ─────────────────────────────────────────────────────────────
page_label(16, 26, "Approved Logo · The mark", "الشعار المعتمد")
add_kicker("Direction C · Final · Approved")
add_h1("The badge.", "الشارة.")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
try:
    run.add_picture("/sessions/dazzling-busy-tesla/mnt/outputs/omlt-logo-approved.png", width=Inches(3.2))
except Exception: pass
p.paragraph_format.space_after = Pt(8)

add_body_ar("شارة دائرية تحوي رمز البيضة في المنتصف، نصًّا مُقوَّسًا حول الأعلى، واسم العلامة OMLT في الأسفل. اعتُمدت رسميًّا كهوية وحيدة للمشروع.")
add_para("File: OMLT-Logo-Layered-HiRes · Approved · Final", size=9, color=GRAY, font="Space Grotesk", rtl=False, space_before=6)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 16 · LOGO ANATOMY
# ─────────────────────────────────────────────────────────────
page_label(17, 26, "Approved Logo · Anatomy", "تشريح الشعار")
add_kicker("Composition")
add_h1("Six layers, one signal.", "ست طبقات، إشارة واحدة.")

anatomy = [
    ("01 · FRAME", "Outer frame", "دائرة معبّأة قرميدية بحدّ كحلي."),
    ("02 · ARC", "Arc text", "“FRESH · DAILY · SINCE 2026” بخط Space Grotesk."),
    ("03 · EGG", "Egg shape", "بيضة عضوية كريمية بحدّ كحلي."),
    ("04 · YOLK", "Yolk #E6A532", "صفار باللون #E6A532 — نقطة التركيز."),
    ("05 · WORDMARK", "OMLT", "OMLT بخط Frank Ruhl Libre 900، كريمي."),
    ("06 · ANCHOR", "Hairline anchor", "زخرفة رفيعة مع نقطة صفراء صغيرة في منتصفها."),
]
tbl = doc.add_table(rows=len(anatomy), cols=3)
tbl.autofit = False
for r in tbl.rows:
    r.cells[0].width = Cm(4); r.cells[1].width = Cm(5); r.cells[2].width = Cm(12)
for code, en, ar in anatomy:
    row = tbl.rows[anatomy.index((code, en, ar))]
    p = row.cells[0].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(code); r.font.name = "Space Grotesk"; r.bold = True; r.font.color.rgb = TERRA; r.font.size = Pt(10)
    p = row.cells[1].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(en); r.font.name = "Frank Ruhl Libre"; r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
    p = row.cells[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p)
    r = p.add_run(ar); r.font.name = "Tajawal"; r.font.size = Pt(11); r.font.color.rgb = NAVY
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 17 · LOGO USAGE RULES
# ─────────────────────────────────────────────────────────────
page_label(18, 26, "Approved Logo · Usage rules", "قواعد الاستخدام")
add_kicker("Do · Don't")
add_h1("A logo doesn't negotiate.", "الشعار لا يُفاوَض.")

HERB_OK = RGBColor(0x4F, 0x80, 0x66)
DEEP_NO = RGBColor(0x9C, 0x35, 0x20)

add_para("✓  CORRECT USAGE", bold=True, size=10, color=HERB_OK, font="Space Grotesk", rtl=False, space_before=10, space_after=4)
do_rules = [
    ("Clear space preserved",       "مساحة آمنة محفوظة حول الشعار"),
    ("Approved background",         "خلفية معتمدة (قرميدي / كحلي / كريمي)"),
    ("Minimum size 32 mm",          "الحجم الأدنى ٣٢ ملم في الطباعة"),
    ("Complete and centered",       "كاملًا في مكانه — لا قصّ ولا حذف"),
]
for en, ar in do_rules:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("✓  "); r.font.name = "Space Grotesk"; r.bold = True; r.font.size = Pt(11); r.font.color.rgb = HERB_OK
    r = p.add_run(en + "   —   "); r.font.name = "Frank Ruhl Libre"; r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
    r = p.add_run(ar); r.font.name = "Tajawal"; r.font.size = Pt(11); r.font.color.rgb = GRAY

add_para("✕  NEVER DO", bold=True, size=10, color=DEEP_NO, font="Space Grotesk", rtl=False, space_before=14, space_after=4)
no_rules = [
    ("Stretch or squash",           "تمديد أو ضغط الشعار"),
    ("Rotation or tilt",            "إمالة أو دوران"),
    ("Color modification",          "تعديل الألوان"),
    ("Off-palette background",      "خلفية خارج لوحة العلامة"),
]
for en, ar in no_rules:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("✕  "); r.font.name = "Space Grotesk"; r.bold = True; r.font.size = Pt(11); r.font.color.rgb = DEEP_NO
    r = p.add_run(en + "   —   "); r.font.name = "Frank Ruhl Libre"; r.bold = True; r.font.size = Pt(12); r.font.color.rgb = NAVY
    r = p.add_run(ar); r.font.name = "Tajawal"; r.font.size = Pt(11); r.font.color.rgb = GRAY
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 18 · COLORS — PALETTE
# ─────────────────────────────────────────────────────────────
page_label(19, 26, "Brand colors · Palette", "لوحة الألوان")
add_kicker("Six colors, fixed ratios")
add_h1("Six colors. Locked ratios.", "ست ألوان بنسب ثابتة.")

palette = [
    ("Terracotta", "قرميدي", "#C8442C", "50%", "C8442C", "F6ECD6"),
    ("Cream",      "كريمي",  "#F6ECD6", "25%", "F6ECD6", "1B2235"),
    ("Navy",       "كحلي",   "#1B2235", "12%", "1B2235", "F6ECD6"),
    ("Yolk",       "صفار",   "#E6A532", "8%",  "E6A532", "1B2235"),
    ("Herb",       "عشبي",   "#7BA284", "3%",  "7BA284", "F6ECD6"),
    ("Deep",       "عميق",   "#9C3520", "2%",  "9C3520", "F6ECD6"),
]
tbl = doc.add_table(rows=len(palette), cols=5)
tbl.autofit = False
widths = [Cm(2.5), Cm(4), Cm(4), Cm(3.5), Cm(3)]
for r in tbl.rows:
    for i, w in enumerate(widths): r.cells[i].width = w
for nm, ar, hex_label, ratio, bg, fg in palette:
    row = tbl.rows[palette.index((nm, ar, hex_label, ratio, bg, fg))]
    shade_cell(row.cells[0], bg)
    # Force a non-empty cell so shading is preserved
    row.cells[0].paragraphs[0].add_run("  ")
    p = row.cells[1].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(nm); r.font.name = "Frank Ruhl Libre"; r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
    p = row.cells[2].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p)
    r = p.add_run(ar); r.font.name = "Tajawal"; r.font.size = Pt(12); r.font.color.rgb = NAVY
    p = row.cells[3].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(hex_label); r.font.name = "Space Grotesk"; r.font.size = Pt(10); r.font.color.rgb = GRAY
    p = row.cells[4].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(ratio); r.font.name = "Frank Ruhl Libre"; r.bold = True; r.font.size = Pt(14); r.font.color.rgb = TERRA

add_para("النسب ثابتة في كل التطبيقات: من اللافتة إلى القائمة إلى التغليف إلى الموقع.",
         size=10, color=GRAY, font="Tajawal", align=WD_ALIGN_PARAGRAPH.RIGHT, space_before=12)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 19 · COLORS — YOLK RULE
# ─────────────────────────────────────────────────────────────
page_label(20, 26, "Brand colors · The yolk rule", "قاعدة الصفار")
add_kicker("One yolk per design")
add_h1("One drop of yolk. Never a sea.", "قطرة صفار واحدة، لا بحرٌ منه.")
add_body_ar("اللون الصفار يظهر لمسة واحدة فقط في كل تصميم — لا اثنتين، ولا يُستخدم كخلفية أبدًا. هذا ما يجعله توقيعًا، لا حشوًا.")
add_body_ar("النسب ثابتة في كل التطبيقات. زيادة لون عن نسبته المحدّدة تُخلّ بشخصية العلامة.")

add_divider()
add_para("✓  CORRECT — نقطة صفار واحدة فقط كتوقيع.", bold=True, size=11, color=RGBColor(0x4F, 0x80, 0x66), font="Frank Ruhl Libre", rtl=False)
add_para("✕  INCORRECT — الصفار خلفيةً، أو الصفار مكرّر في نفس التصميم.", bold=True, size=11, color=RGBColor(0x9C, 0x35, 0x20), font="Frank Ruhl Libre", rtl=False)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 20 · TYPE — FRANK RUHL LIBRE
# ─────────────────────────────────────────────────────────────
page_label(21, 26, "Typography · Frank Ruhl Libre", "الخط الأول")
add_kicker("Display · Headlines")
add_h1("Frank Ruhl Libre.", "الخط الأول — للعناوين.")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("A breakfast kitchen."); r.font.name = "Frank Ruhl Libre"; r.bold = True; r.font.size = Pt(40); r.font.color.rgb = NAVY

add_body_ar("الاستخدام — العناوين الرئيسية، اسم العلامة، الجمل التحريرية، الاقتباسات.")
add_body_ar("الأوزان — 400 عادي، 500 متوسط، 700 عريض، 900 للعناوين.")
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 21 · TYPE — SPACE GROTESK + TAJAWAL
# ─────────────────────────────────────────────────────────────
page_label(22, 26, "Typography · Space Grotesk & Tajawal", "الخطان الثاني والثالث")
add_kicker("UI · Body · Arabic")
add_h2("Space Grotesk & Tajawal.")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = p.add_run("Mushroom omelette · 38 SAR"); r.font.name = "Space Grotesk"; r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY
add_body_ar("Space Grotesk — للنصوص الإنجليزية والقائمة والأرقام وواجهات الموقع. أوزان: 400، 500، 600، 700.")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT; set_rtl(p); p.paragraph_format.space_before = Pt(10)
r = p.add_run("تجوال · مطبخ فطور صغير"); r.font.name = "Tajawal"; r.bold = True; r.font.size = Pt(28); r.font.color.rgb = NAVY
add_body_ar("Tajawal — كل النصوص العربية بدون استثناء. أوزان: 500 متن، 700 تأكيد، 900 عناوين فقط. ارتفاع السطر 1.85.")

add_divider()
add_body_ar("قواعد عامة — العربية أساس، الإنجليزية مساندة. لا italic للعربية. لا خط رابع.")
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 22 · VISUAL IDENTITY — PERSONALITY
# ─────────────────────────────────────────────────────────────
page_label(23, 26, "Visual identity · Personality", "شخصية العلامة")
add_kicker("Brand personality")
add_h1("Warm · Clear · Brief · Signed.", "دافئة · واضحة · مختصرة · موقّعة.")

pers = [
    ("Warm · دافئة", "ألوان طبيعية، إضاءة هادئة، إحساس بمكان مألوف وليس رسمي."),
    ("Clear · واضحة", "بدون مبالغات بصرية أو زخرفة. كل عنصر له وظيفة وسبب وجود."),
    ("Brief · مختصرة", "لغة قصيرة، بدون كلمات كثيرة. ثقة في فهم العميل دون شرح طويل."),
    ("Signed · موقّعة", "ظهور واحد للشارة في كل سطح — على الكوب، على اللافتة، على الإيصال."),
]
for en, ar in pers:
    add_para(en, bold=True, size=12, color=TERRA, font="Frank Ruhl Libre", rtl=False, space_before=6)
    add_body_ar(ar, size=11)
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 23 · SENSORY IDENTITY
# ─────────────────────────────────────────────────────────────
page_label(24, 26, "Sensory identity · Four senses", "الهوية الحسّية")
add_kicker("Four sensory gates")
add_h1("Sight · Smell · Touch · Warmth.", "البصر · الشمّ · اللمس · الحرارة.")

senses = [
    ("Sight · البصر", "إضاءة دافئة، مطبخ مفتوح ومرئي، مساحة منظَّمة وبسيطة. لا شاشات ولا مؤثّرات بصرية."),
    ("Smell · الشمّ", "روائح الخبز والبيض والقهوة من المطبخ مباشرة، وليست من معطّرات صناعية."),
    ("Touch · اللمس", "أكواب وأطباق وأدوات بجودة جيّدة، ليست فاخرة لكنها متينة ومريحة في الاستخدام اليومي."),
    ("Warmth · الحرارة", "إضاءة دافئة بحدود 2700K. لا ضوء أبيض بارد. حرارة المكان مريحة طوال الصباح."),
]
for en, ar in senses:
    add_para(en, bold=True, size=12, color=TERRA, font="Frank Ruhl Libre", rtl=False, space_before=6)
    add_body_ar(ar, size=11)

add_divider()
add_body_ar("تجربة العميل المعتاد — معرفة العملاء بشكل طبيعي، بدون استعراض. الموظف يتذكّر الوجوه والطلبات بشكل عمليّ يُسرّع الخدمة.")
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 24 · AUDIO IDENTITY
# ─────────────────────────────────────────────────────────────
page_label(25, 26, "Audio identity · Voice of the room", "الهوية الصوتية")
add_kicker("Sound of the morning")
add_h1("Let the kitchen speak.", "دع المطبخ يتكلّم.")
add_body_ar("خلال ساعات الصباح، تكون أصوات المطبخ — تحضير البيض، حركة المقلاة، تجهيز القهوة — جزءًا واضحًا من تجربة المكان. ويمكن استخدام موسيقى هادئة منخفضة عند الحاجة، دون أن تطغى على أصوات المطبخ أو حديث العملاء.")

audio_rules = [
    "أصوات المطبخ هي العنصر الأساسي في الجو الصوتي",
    "موسيقى هادئة منخفضة عند الحاجة — لا تُستخدم كملء فراغ",
    "لا موسيقى صاخبة أو أنواع لا تناسب جو الصباح",
    "معالجة صوتية للمكان لتقليل الصدى والضوضاء",
    "التحيات طبيعية ومباشرة، بدون نصوص محفوظة",
]
for b in audio_rules: add_bullet_ar(b)

add_divider()
add_h3("نبرة التواصل المكتوب · Written tone")
add_body_ar("اللغة في المنشورات والقائمة واللافتات: قصيرة، واضحة، بدون مبالغات تسويقية. نفس النبرة في كل القنوات.")
page_break()

# ─────────────────────────────────────────────────────────────
# PAGE 25 · CLOSING
# ─────────────────────────────────────────────────────────────
page_label(26, 26, "End of archive", "ختام الأرشيف")

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run()
try:
    run.add_picture("/sessions/dazzling-busy-tesla/mnt/outputs/omlt-logo-approved.png", width=Inches(2.4))
except Exception: pass
p.paragraph_format.space_after = Pt(16)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("The OMLT promise"); r.font.name = "Space Grotesk"; r.bold = True; r.font.size = Pt(10); r.font.color.rgb = TERRA
p.paragraph_format.space_after = Pt(6)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("A clear and simple morning.")
r.font.name = "Frank Ruhl Libre"; r.bold = True; r.font.size = Pt(36); r.font.color.rgb = NAVY
p.paragraph_format.space_after = Pt(8)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; set_rtl(p)
r = p.add_run("بداية صباح واضحة وبسيطة.")
r.font.name = "Tajawal"; r.bold = True; r.font.size = Pt(20); r.font.color.rgb = TERRA
p.paragraph_format.space_after = Pt(14)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("OMLT · MASTER ARCHIVE · ABHA · KSA")
r.font.name = "Space Grotesk"; r.font.size = Pt(9); r.font.color.rgb = GRAY

out = "/sessions/dazzling-busy-tesla/mnt/outputs/OMLT-Master-Presentation.docx"
doc.save(out)
print("OK:", out)
